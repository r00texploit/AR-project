from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs" / "ar"
OUT_DIR = DOCS_DIR / "docx"

FONT = "Arial"


def clean_inline(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = text.replace("\\|", "|")
    text = text.replace("`", "")
    return text


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_rtl(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_cell_text(cell, text, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    run = p.add_run(clean_inline(text.strip()))
    set_run_font(run, 10, bold=bold)
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag, value in {
        "w:top": "80",
        "w:bottom": "80",
        "w:start": "120",
        "w:end": "120",
    }.items():
        tc_mar = tc_pr.find(qn("w:tcMar"))
        if tc_mar is None:
            tc_mar = OxmlElement("w:tcMar")
            tc_pr.append(tc_mar)
        el = tc_mar.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            tc_mar.append(el)
        el.set(qn("w:w"), value)
        el.set(qn("w:type"), "dxa")


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color in [
        ("Heading 1", 20, "1F4D78"),
        ("Heading 2", 16, "2E74B5"),
        ("Heading 3", 13, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    set_paragraph_rtl(p)
    run = p.add_run(clean_inline(text))
    set_run_font(run, {1: 20, 2: 16, 3: 13}.get(level, 12), bold=True,
                 color="1F4D78" if level != 2 else "2E74B5")


def add_body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    text = clean_inline(text)
    for part in re.split(r"(`[^`]+`)", text):
        if not part:
            continue
        run = p.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        set_run_font(run, 11, bold=False)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Courier New"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
            run._element.rPr.rFonts.set(qn("w:cs"), "Courier New")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_rtl(p)
    run = p.add_run(clean_inline(text))
    set_run_font(run, 11)


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_rtl(p)
    run = p.add_run(clean_inline(text))
    set_run_font(run, 11)


def add_code(doc, lines):
    p = doc.add_paragraph()
    set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.left_indent = Inches(0.2)
    for i, line in enumerate(lines):
        if i:
            p.add_run("\n")
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run._element.rPr.rFonts.set(qn("w:cs"), "Courier New")
        run.font.size = Pt(9.5)


def add_table(doc, rows):
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(width):
            text = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), text, bold=(r_idx == 0))
            if r_idx == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F4F7")
                table.cell(r_idx, c_idx)._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def convert_md(md_path, docx_path):
    doc = Document()
    style_document(doc)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if match:
            add_heading(doc, match.group(2), len(match.group(1)))
        elif re.match(r"^[-*]\s+", stripped):
            add_bullet(doc, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+\.\s+", stripped):
            add_numbered(doc, re.sub(r"^\d+\.\s+", "", stripped))
        else:
            add_body(doc, stripped)
        i += 1

    if in_code and code_lines:
        add_code(doc, code_lines)

    footer = doc.sections[0].footer.paragraphs[0]
    set_paragraph_rtl(footer)
    run = footer.add_run("AR Education MVP - توثيق عربي")
    set_run_font(run, 9, color="666666")

    doc.save(docx_path)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = [
        DOCS_DIR / "README.md",
        DOCS_DIR / "01-نظرة-عامة.md",
        DOCS_DIR / "02-تشغيل-التطبيق.md",
        DOCS_DIR / "03-بناء-ملف-apk.md",
        DOCS_DIR / "04-المعمارية-والبيانات.md",
    ]
    for md_path in files:
        out_name = md_path.with_suffix(".docx").name
        convert_md(md_path, OUT_DIR / out_name)
        print(OUT_DIR / out_name)


if __name__ == "__main__":
    main()
